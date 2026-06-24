import os
import io
import json
import uuid
import shutil
import jwt
import re
from datetime import datetime, timedelta, timezone
from typing import List, Optional

from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, Request
from fastapi.security import OAuth2PasswordBearer
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.errors import RateLimitExceeded
from slowapi.util import get_remote_address

from docx import Document

# Import Database tools
from databaseFile import engine, Base, get_db
import models

# Import your AI engine
from ai_engine import ask_laguna

JWT_SECRET = os.getenv("JWT_SECRET", "change_this_secret")
ALGORITHM = "HS256"
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login")


def verify_token(token: str = Depends(oauth2_scheme)):
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[ALGORITHM])
        return payload.get("sub")
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")


limiter = Limiter(key_func=get_remote_address)

Base.metadata.create_all(bind=engine)
app = FastAPI()
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


class ChatMessageDef(BaseModel):
    role: str
    content: str


class ChatRequest(BaseModel):
    message: str
    filename: str = "bim_data.txt"
    history: List[ChatMessageDef] = []
    chat_id: Optional[int] = None


class CommentRequest(BaseModel):
    chat_id: int
    author: str
    content: str


@app.post("/login")
@limiter.limit("10/minute")
def login_placeholder(request: Request):
    expiration_time = datetime.now(timezone.utc) + timedelta(hours=24)
    token_data = {"sub": "test_engineer", "exp": expiration_time}
    encoded_jwt = jwt.encode(token_data, JWT_SECRET, algorithm=ALGORITHM)
    return {"access_token": encoded_jwt, "token_type": "bearer"}


@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    if not file.filename.endswith(('.pdf', '.txt')):
        raise HTTPException(status_code=400, detail="Only .pdf and .txt files are allowed.")
    file_location = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_location, "wb+") as file_object:
        shutil.copyfileobj(file.file, file_object)
    return {"message": "Success", "filename": file.filename, "path": file_location}


@app.get("/documents")
async def list_documents():
    files = [f for f in os.listdir(UPLOAD_DIR) if f.endswith(('.pdf', '.txt'))]
    return {"documents": files}


@app.get("/download/{filename}")
async def download_file(filename: str):
    filepath = os.path.join(UPLOAD_DIR, filename)
    if os.path.exists(filepath):
        return FileResponse(filepath, filename=filename)
    raise HTTPException(status_code=404, detail="File not found")


@app.post("/chat")
@limiter.limit("20/minute")
async def chat_endpoint(request: Request, chat_request: ChatRequest, db: Session = Depends(get_db)):
    if not chat_request.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty.")

    current_chat_id = chat_request.chat_id
    chat_record = db.query(models.Chat).filter(models.Chat.id == current_chat_id).first()

    if not chat_record:
        new_chat = models.Chat(id=current_chat_id, title=chat_request.message[:30])
        db.add(new_chat)
        db.commit()

    user_msg = models.Message(chat_id=current_chat_id, role="user", content=chat_request.message)
    db.add(user_msg)
    db.commit()

    file_path = os.path.join(UPLOAD_DIR,
                             chat_request.filename) if chat_request.filename != "bim_data.txt" else "bim_data.txt"
    dict_history = [{"role": msg.role, "content": msg.content} for msg in chat_request.history]

    # 1. Fetch AI Response
    ai_response = ask_laguna(chat_request.message, file_path, dict_history)

    chart_data = None
    document_filename = None

    # 2. Extract Intent: Chart
    chart_match = re.search(r'<chart>(.*?)</chart>', ai_response, re.DOTALL)
    if chart_match:
        try:
            chart_data = json.loads(chart_match.group(1))
            ai_response = re.sub(r'<chart>.*?</chart>', '', ai_response, flags=re.DOTALL).strip()
        except Exception:
            pass

    # 3. Extract Intent: Document Generation
    doc_match = re.search(r'<document>(.*?)</document>', ai_response, re.DOTALL)
    if doc_match:
        doc_content = doc_match.group(1)
        ai_response = re.sub(r'<document>.*?</document>', '', ai_response, flags=re.DOTALL).strip()

        doc = Document()
        doc.add_heading("BIM Generated Report", level=1)
        for paragraph in doc_content.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)

        filename = f"report_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(UPLOAD_DIR, filename)
        doc.save(filepath)
        document_filename = filename

    # Fallback if empty
    if not ai_response.strip():
        ai_response = "Here is the generated output based on your request:"

    ai_msg = models.Message(chat_id=current_chat_id, role="assistant", content=ai_response)
    db.add(ai_msg)
    db.commit()

    return {
        "response": ai_response,
        "chat_id": current_chat_id,
        "chart_data": chart_data,
        "document_url": f"/download/{document_filename}" if document_filename else None
    }


@app.post("/comments")
async def add_comment(req: CommentRequest, db: Session = Depends(get_db)):
    comment = models.Comment(chat_id=req.chat_id, author=req.author, content=req.content)
    db.add(comment)
    db.commit()
    return {"status": "added"}


@app.get("/comments/{chat_id}")
async def get_comments(chat_id: int, db: Session = Depends(get_db)):
    comments = db.query(models.Comment).filter(models.Comment.chat_id == chat_id).all()
    return [{"author": c.author, "content": c.content, "created_at": c.created_at} for c in comments]